import sys
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor


def validateInput():
    """ Validates format of command line input from the user, prints error,  and
    quits program if the format is incorrect
    Returns:
        int word count in jobTitle, int word count in companyName
    """
    # validating that value in argv[1] is an int specifying number of words in company name

    try:
        wordsInCompanyName = int(sys.argv[1])
    except:
        print("Must supply integer value for words in company name first. Error, quitting\n")
        quit()

    # validating that the int provided in argv[1] is accurate to words in company name

    try:
        wordsInJobTitleArgIndex = 2 + wordsInCompanyName

        wordsInJobTitle = int(sys.argv[wordsInJobTitleArgIndex])

    except:
        print("Incorrect integer value for words in company name provided. Error, quitting\n")
        quit()

    # validate that the ints provided on company name and job title word count are accurate

    argCountProvided= 3 + wordsInJobTitle + wordsInCompanyName

    if argCountProvided != len(sys.argv):
        print("Incorrect integer value provided for words in job title. Error, quitting\n")
        quit()

    return wordsInJobTitle, wordsInCompanyName


def buildDictionaryForDocEdits(wordsInCompanyName, wordsInJobTitle):
    """ Utilizes the sys.argv inputs to fill a dictionary with the values of the
    company name and job title and returns the dictionary
    Args: int wordsInCompanyName, int wordsInJobTitle
    Returns:
        dict values (job title and company name)
    """


    values = {}
    companyName = ""
    jobTitle = ""

    for i in range(1, wordsInCompanyName+1):
        companyName += str(sys.argv[1+i])
        companyName += " "

    companyName.strip()

    #print("Company Name: \n", companyName)


    for i in range(1, wordsInJobTitle+1):
        jobTitle += str(sys.argv[2+i+wordsInCompanyName])
        jobTitle += " "

    jobTitle.strip()

    #print("Job Title: \n", jobTitle)

    aposLocation = determineApostropheLocation(companyName.strip())

    if aposLocation == 1:
        punctuatedCompany = companyName.strip()
        punctuatedCompany += "\'"
    else:
        punctuatedCompany = companyName.strip()
        punctuatedCompany += "\'s"

    values["Company"] = companyName.strip()
    values["Job Title"] = jobTitle.strip()
    values["pluralComp"] = punctuatedCompany.strip()

    return values


def determineApostropheLocation(companyName):
    if companyName[-1] == 's' or companyName[-1] == 'S':
        return 1
    else:
        return 0


def editAndSaveDocxFile(values):
    """ Edits the cover letter adding company name, job title, and querying
    the user for input on why they are interested in the company. It searches
    for keywords 'Company', 'Job Title', and 'pluralCompany' to substitute with the
    user inputted values. Saves Document to the /Downloads folder.
    Args: Dict values with company name and job title
    Returns:
        docx file in Downloads folder with the title: job title cover letter.docx
    """
    # get user input for specifics about their interest in the company
    companyPraise = getInput(values)

    doc = Document('genericCoverLetter.docx')

    # running through document and checking for keys in 'values' dict to be replaced with the value
    for i in values:
        for p in doc.paragraphs:
            if p.text.find(i) >= 0:
                #print('found')
                p.text = p.text.replace(i,values[i])

    # adding userInput, company praise into the cover letter

    paragraphNum = 0
    for paragraph in doc.paragraphs:
        if paragraphNum == 2:
            paragraph.add_run(companyPraise)
        paragraphNum += 1
        print(paragraph.text)

    # changing font of the document

    for paragraph in doc.paragraphs:
        paragraph.style = doc.styles['Normal']
        for run in paragraph.runs:
            run.font.size = Pt(12)
            run.font.name = 'Arial'
            black = RGBColor(0, 0, 0)
            run.font.color.rgb = black

    # save cover letter file to /Downloads folder with company name

    doc.save('..//Downloads/{} Cover Letter.docx'.format(values["Company"]))


def getInput(values):
    userInput = input("I am highly impressed with {} ".format(values["pluralComp"]))
    return userInput


#---------------------------------------------------------------------------------#

if __name__ == "__main__":

    wordsInJobTitle, wordsInCompanyName = validateInput()

    values = buildDictionaryForDocEdits(wordsInCompanyName, wordsInJobTitle)

    editAndSaveDocxFile(values)
